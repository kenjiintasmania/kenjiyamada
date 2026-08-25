# -*- coding: utf-8 -*-
"""tools/make_paper.py ─ 模試データ(mogi/data/*.js)から「紙媒体テスト」のDOCXを作る

使い方:
    python3 tools/make_paper.py <データファイル.js> <出力先ディレクトリ> [名前]

作るもの（2ファイル）:
    ・<名前>_問題.docx   … 生徒に配るもの。解答は載せない
    ・<名前>_解答.docx   … 先生用。放送文（読み上げ台本）＋解答一覧＋配点内訳

紙にするときの約束:
  ・リスニングで「選択肢そのものが放送される」大問（scriptを持たず英文の選択肢を持つ群）は、
    選択肢を生徒用紙に印刷しない。印刷すると読むだけで解けてしまうため。
    その選択肢は先生用の読み上げ台本にまわす。
  ・コース選択（X/Y）のある大問は両方を印刷し、どちらか一方を選ばせる。配点は1コース分で数える。
"""
import json, re, subprocess, sys, os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

KANA = "アイウエオカキク"

# ---------- JS の EXAM を JSON にして取り出す ----------
def load_exam(js_path):
    code = (
        'const fs=require("fs");const w={};'
        'new Function("window",fs.readFileSync(process.argv[1],"utf8"))(w);'
        'process.stdout.write(JSON.stringify(w.EXAM));'
    )
    out = subprocess.run(["node", "-e", code, js_path], capture_output=True, text=True)
    if out.returncode != 0:
        sys.exit("✗ データを読めません: " + out.stderr.strip()[:300])
    return json.loads(out.stdout)

# ---------- HTML → 紙用のテキスト ----------
def strip(html, keep_speaker=True):
    s = str(html or "")
    s = re.sub(r'<div class="scene">(.*?)</div>', r'【\1】', s, flags=re.S)
    s = s.replace("</span><span class=\"sp\">", "\n")
    if keep_speaker:
        s = re.sub(r'<span class="who"[^>]*>(.*?)</span>', r'\1 ', s, flags=re.S)
    s = re.sub(r'<br\s*/?>', "\n", s)
    s = re.sub(r'</(p|div|tr|h4)>', "\n", s)
    s = re.sub(r'<[^>]+>', "", s)
    s = s.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    s = re.sub(r'[ \t　]+', " ", s)
    s = re.sub(r'\n{3,}', "\n\n", s)
    return s.strip()

def html_tables(html):
    """flyer 内の <table> を [[行, 行...], ...] で返す"""
    out = []
    for tb in re.findall(r'<table.*?</table>', str(html or ""), flags=re.S):
        rows = []
        for tr in re.findall(r'<tr.*?</tr>', tb, flags=re.S):
            cells = [strip(c) for c in re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', tr, flags=re.S)]
            if any(cells):
                rows.append(cells)
        if rows:
            out.append(rows)
    return out

def html_wo_tables(html):
    return strip(re.sub(r'<table.*?</table>', "", str(html or ""), flags=re.S))

# ---------- 書式 ----------
# 使える本文幅（A4 21.0cm − 左右余白）。表がはみ出さないように使う。
USABLE_W = Cm(21.0 - 1.7 * 2)

def setup(doc):
    st = doc.styles["Normal"]
    st.font.name = "游明朝"
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
    for s in doc.sections:
        # ★A4を明示する。python-docx の既定テンプレートは米国レター(21.6×27.9cm)で、
        #   指定を省くとA4プリンタで用紙不一致になり、縮小されて隅に印刷される。
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = s.bottom_margin = Cm(1.6)
        s.left_margin = s.right_margin = Cm(1.7)
        s.header_distance = Cm(1.0)
        s.footer_distance = Cm(1.0)

def para(doc, text="", size=10.5, bold=False, align=None, before=0, after=2,
         italic=False, color=None, en=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    for i, line in enumerate(str(text).split("\n")):
        r = p.add_run(("" if i == 0 else "\n") + line)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
        name = "Century" if en else "游明朝"
        r.font.name = name
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
    return p

def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    bd = pPr.makeelement(qn("w:pBdr"), {})
    bt = bd.makeelement(qn("w:bottom"), {})
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "888888")):
        bt.set(qn(k), v)
    bd.append(bt); pPr.append(bd)

def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False                      # 本文幅を超えて右にはみ出すのを防ぐ
    for c in t.columns:
        c.width = Emu(int(USABLE_W / cols))
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j); c.text = ""
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(cell); r.font.size = Pt(9)
            r.font.name = "游明朝"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
            if i == 0: r.bold = True
    return t

# ---------- 設問の判定 ----------
def is_listening_section(sec):
    """リスニングの大問か（放送台本を持つ群がある／表題にリスニングを含む）"""
    if re.search(r'リスニング|放送', strip(sec.get("title", ""))):
        return True
    return any(g.get("script") for g in sec.get("groups", []))

def is_audio_choice(sec, group, item):
    """選択肢そのものが放送される＝生徒用紙に印刷してはいけない設問か。
       読解の大問で誤判定すると選択肢が消えてしまうため、リスニングの大問に限定する。"""
    if not is_listening_section(sec):
        return False
    if group.get("script"):
        return False
    if item.get("type") != "mcq":
        return False
    chs = item.get("choices") or []
    # 選択肢が英文で、設問文が「絵：」などの日本語の説明のとき
    return bool(chs) and all(re.search(r'[A-Za-z]', strip(c)) for c in chs) \
           and not re.search(r'[A-Za-z]{3}', strip(item.get("stem", "")))

def items_of(sec):
    """(コース名, item, group) の並びを返す。コース制の大問にも対応。"""
    out = []
    if sec.get("courses"):
        for c in sec["courses"]:
            for it in c["items"]:
                out.append((c.get("name", ""), it, {}))
    else:
        for g in sec.get("groups", []):
            for it in g.get("items", []):
                out.append(("", it, g))
    return out

def section_points(sec):
    """生徒が実際に取れる満点。コース制は1コース分だけ数える。"""
    if sec.get("courses"):
        return max(sum(i.get("pt", 0) for i in c["items"]) for c in sec["courses"])
    return sum(i.get("pt", 0) for _, i, _ in items_of(sec))

def blanks(n=1, w=26):
    return "\n".join(["　" + "＿" * w for _ in range(n)])

# ---------- ① 問題用紙 ----------
def build_question(exam, path, title):
    doc = Document(); setup(doc)
    total = sum(section_points(s) for s in exam["sections"])
    para(doc, title, 16, True, WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "（{}点満点）".format(total), 9.5, False, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
    for j, lab in enumerate(["学年", "組", "番号", "名前"]):
        c = t.cell(0, j); c.text = ""
        r = c.paragraphs[0].add_run(lab + "　　　　　")
        r.font.size = Pt(10); r.font.name = "游明朝"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")
    para(doc, "", 6, after=4)

    for sec in exam["sections"]:
        rule(doc)
        para(doc, "大問{}　{}　（{}点）".format(sec["no"], strip(sec.get("title", "")),
             section_points(sec)), 12, True, before=6, after=3)
        if sec.get("lead"):
            lead = strip(sec["lead"])
            # アプリ用の文言（画面で放送文を読む前提）は紙では成り立たないので置きかえる
            lead = lead.replace("放送文を読んで、", "放送を聞いて、")
            lead = re.sub(r'（実際の試験では音声が流れます）', "", lead).strip()
            para(doc, lead, 10, after=3)

        if sec.get("courses"):
            para(doc, "※ XコースとYコースのどちらか一方を選び、選んだコースに○をつけて答えなさい。"
                      "（　X　・　Y　）", 10, True, after=4)
            for c in sec["courses"]:
                para(doc, "【{}】".format(c.get("name", "")), 10.5, True, before=4, after=2)
                for it in c["items"]:
                    render_item(doc, it, {}, printable=True)
            continue

        for g in sec.get("groups", []):
            if g.get("intro"):
                para(doc, strip(g["intro"]), 10, before=4, after=3)
            if g.get("flyer"):
                head = html_wo_tables(g["flyer"])
                if head: para(doc, head, 10, True, after=2)
                for rows in html_tables(g["flyer"]):
                    add_table(doc, rows); para(doc, "", 6, after=2)
            if g.get("passage"):
                para(doc, strip(g["passage"]), 10.5, after=3, en=True)
            if g.get("note"):
                para(doc, strip(g["note"]), 9, after=3, italic=True)
            for it in g.get("items", []):
                render_item(doc, it, g, printable=not is_audio_choice(sec, g, it))
    doc.save(path)
    return total

def render_item(doc, it, group, printable=True):
    lab = it.get("label", "")
    head = "{}　（{}点）".format(lab, it.get("pt", 0)) if lab else "（{}点）".format(it.get("pt", 0))
    stem = strip(it.get("stem", ""))
    para(doc, head + ("　" + stem if stem else ""), 10.5, before=3, after=2)
    ty = it.get("type")
    if ty in ("mcq", "bankpick"):
        arr = it.get("choices") if ty == "mcq" else it.get("bank")
        if printable:
            for i, c in enumerate(arr or []):
                para(doc, "　{}　{}".format(KANA[i], strip(c)), 10, after=1, en=True)
            para(doc, "　答え（　　　　）", 10, after=2)
        else:
            para(doc, "　※ 選択肢は放送されます。　答え（　　　　）", 10, after=2)
    elif ty == "fill":
        para(doc, blanks(1), 10.5, after=2)
    elif ty == "wordorder":
        para(doc, blanks(1), 10.5, after=2)

# ---------- ② 解答・配点（先生用） ----------
def build_answer(exam, path, title):
    doc = Document(); setup(doc)
    total = sum(section_points(s) for s in exam["sections"])
    para(doc, title + "　解答・配点（先生用）", 15, True, WD_ALIGN_PARAGRAPH.CENTER, after=6)

    para(doc, "■ 放送文（読み上げ台本）", 12, True, before=4, after=3)
    read = 0
    for sec in exam["sections"]:
        for g in sec.get("groups", []):
            if g.get("script"):
                read += 1
                para(doc, strip(g.get("intro", "")), 9.5, True, after=2)
                para(doc, strip(g["script"]), 10.5, after=4, en=True)
            else:
                auds = [it for it in g.get("items", []) if is_audio_choice(sec, g, it)]
                if auds:
                    read += 1
                    para(doc, strip(g.get("intro", "")) + "（選択肢を読み上げてください）", 9.5, True, after=2)
                    for it in auds:
                        para(doc, "{}　{}".format(it.get("label", ""), strip(it.get("stem", ""))), 9.5, after=1)
                        for i, c in enumerate(it.get("choices") or []):
                            para(doc, "　　{}　{}".format(KANA[i], strip(c)), 10, after=0, en=True)
                        para(doc, "", 6, after=2)
    if not read:
        para(doc, "（放送文はありません）", 10, after=3)

    para(doc, "■ 解答一覧", 12, True, before=6, after=3)
    rows = [["大問", "問", "正解", "配点"]]
    for sec in exam["sections"]:
        for cname, it, _ in items_of(sec):
            ty = it.get("type")
            if ty == "mcq":
                ans = "{}（{}）".format(KANA[it.get("answer", 0)], strip((it.get("choices") or ["?"])[it.get("answer", 0)]))
            elif ty == "bankpick":
                ans = "{}（{}）".format(KANA[it.get("answer", 0)], strip((it.get("bank") or ["?"])[it.get("answer", 0)]))
            elif ty == "fill":
                ans = " / ".join(strip(a) for a in (it.get("answers") or []))
            elif ty == "wordorder":
                ans = strip(it.get("display") or it.get("answer", ""))
            else:
                ans = ""
            rows.append(["大問{}{}".format(sec["no"], ("・" + cname) if cname else ""),
                         it.get("label", ""), ans, str(it.get("pt", 0))])
    add_table(doc, rows)

    para(doc, "■ 配点の内訳", 12, True, before=8, after=3)
    br = [["大問", "満点"]]
    for sec in exam["sections"]:
        br.append(["大問{}　{}".format(sec["no"], strip(sec.get("title", ""))[:26]), str(section_points(sec))])
    br.append(["合計", str(total)])
    add_table(doc, br)
    para(doc, "※ コース選択のある大問は、1コース分を満点として数えています。", 9, before=4, italic=True)
    doc.save(path)
    return total


# ---------- ③ 解答用紙（A4 1枚・採点／スキャン読み取り前提） ----------
COLS = 10  # 10列の格子。欄の位置が予測できるので、スキャンからの読み取りがしやすい

def _cell_text(cell, text, size=9, bold=False, align=None, en=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    r.font.name = "Century" if en else "游明朝"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "游明朝")

def _anchors(doc, title=None):
    """四隅の位置合わせ用マーク。スキャン画像の傾き・切れの判定に使う。"""
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    for i, wcm in enumerate((1.2, 15.2, 1.2)):
        t.columns[i].width = Cm(wcm)
    _cell_text(t.cell(0, 0), "■", 11, True)
    _cell_text(t.cell(0, 1), title or "", 13, True, WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(t.cell(0, 2), "■", 11, True, WD_ALIGN_PARAGRAPH.RIGHT)
    return t

def _shade(cell, color="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = tcPr.makeelement(qn("w:shd"), {})
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), color)
    tcPr.append(sh)


def _grid(doc, nrows):
    t = doc.add_table(rows=nrows, cols=COLS)
    t.style = "Table Grid"; t.autofit = False
    for c in t.columns:
        c.width = Emu(int(USABLE_W / COLS))
    return t

def _span(t, r, c0, c1):
    return t.cell(r, c0).merge(t.cell(r, c1))

def answer_fields(exam):
    """解答用紙に並べる記入欄。(大問, コース名, item, 種別) を返す。"""
    out = []
    for sec in exam["sections"]:
        for cname, it, _ in items_of(sec):
            ty = it.get("type")
            if ty in ("mcq", "bankpick"):
                kind = "sel"
            elif ty == "wordorder":
                kind = "long"
            else:
                first = strip((it.get("answers") or [""])[0])
                kind = "long" if len(first.split(" ")) > 2 else "word"
            out.append((sec, cname, it, kind))
    return out

def group_prefix(g):
    """群の見出しから (1) などの小問番号を拾う。大問1の①が3回出るような重複を防ぐ。"""
    m = re.match(r'^\((\d+)\)', strip(g.get("intro", "")))
    return "({})".format(m.group(1)) if m else ""

def kind_of(it):
    ty = it.get("type")
    if ty in ("mcq", "bankpick"):
        return "sel"
    if ty == "wordorder":
        return "long"
    first = strip((it.get("answers") or [""])[0])
    return "long" if len(first.split(" ")) > 2 else "word"

def label_of(g, it, pre):
    lab = str(it.get("label", "") or "")
    if lab.startswith("(") and pre and lab.startswith(pre):
        return lab                      # 既に (1)① の形になっている
    return (pre + lab) if lab else (pre or "並")


def build_sheet(exam, path, title):
    doc = Document(); setup(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.2)
        s.left_margin = s.right_margin = Cm(1.7)
    total = sum(section_points(s) for s in exam["sections"])

    _anchors(doc, title + "　解答用紙")
    para(doc, "", 4, after=2)

    # 受験者情報と合計欄
    t = _grid(doc, 1)
    _cell_text(_span(t, 0, 0, 0), "学年", 9, True)
    _cell_text(_span(t, 0, 1, 1), "", 9)
    _cell_text(_span(t, 0, 2, 2), "組", 9, True)
    _cell_text(_span(t, 0, 3, 3), "", 9)
    _cell_text(_span(t, 0, 4, 4), "番号", 9, True)
    _cell_text(_span(t, 0, 5, 5), "", 9)
    _cell_text(_span(t, 0, 6, 6), "名前", 9, True)
    _cell_text(_span(t, 0, 7, 8), "", 9)
    _cell_text(_span(t, 0, 9, 9), "合計\n/{}".format(total), 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    t.rows[0].height = Cm(1.1)
    para(doc, "※ 解答は下の枠の中に書くこと。枠からはみ出すと採点できません。", 8.5, after=3)

    for sec in exam["sections"]:
        # 記入欄を「出題順」に並べる。順を入れ替えると採点時に問と欄が対応しなくなる。
        seq = []
        if sec.get("courses"):
            # X/Yは同じ問番号なので欄は1組でよい（生徒はどちらか一方だけ答える）
            for it in sec["courses"][0]["items"]:
                seq.append((label_of({}, it, ""), kind_of(it)))
        else:
            for g in sec.get("groups", []):
                pre = group_prefix(g)
                for it in g.get("items", []):
                    seq.append((label_of(g, it, pre), kind_of(it)))

        # 行を組み立てる：選択式は連続する分だけ5個ずつ詰め、他は1〜2個で1行
        lines = []   # ("sel",[..]) / ("word",[..]) / ("long",[lab])
        buf_s, buf_w = [], []
        def flush():
            while buf_s:
                lines.append(("sel", buf_s[:5])); del buf_s[:5]
            while buf_w:
                lines.append(("word", buf_w[:2])); del buf_w[:2]
        for lab, kind in seq:
            if kind == "sel":
                if buf_w: flush()
                buf_s.append(lab)
                if len(buf_s) == 5: lines.append(("sel", buf_s[:])); buf_s.clear()
            elif kind == "word":
                if buf_s: flush()
                buf_w.append(lab)
                if len(buf_w) == 2: lines.append(("word", buf_w[:])); buf_w.clear()
            else:
                flush(); lines.append(("long", [lab]))
        flush()

        rows = 1 + len(lines) + (1 if sec.get("courses") else 0)
        t = _grid(doc, rows)
        r = 0
        _cell_text(_span(t, r, 0, 7), "大問{}　{}".format(sec["no"], strip(sec.get("title", ""))[:28]), 9.5, True)
        _cell_text(_span(t, r, 8, 9), "小計 　/{}".format(section_points(sec)), 9, True,
                   WD_ALIGN_PARAGRAPH.CENTER)
        t.rows[r].height = Cm(0.52); r += 1

        if sec.get("courses"):
            _cell_text(_span(t, r, 0, 9),
                       "選んだコースに○　（　　X コース　　・　　Y コース　　）　※どちらか一方だけ答える",
                       9, True)
            t.rows[r].height = Cm(0.55); r += 1

        for kind, labs in lines:
            if kind == "sel":
                for j, lab in enumerate(labs):
                    _cell_text(t.cell(r, j * 2), lab, 8, True, WD_ALIGN_PARAGRAPH.CENTER)
                    _cell_text(t.cell(r, j * 2 + 1), "", 11, align=WD_ALIGN_PARAGRAPH.CENTER)
                used = len(labs) * 2
                if used < COLS:                      # 余った枠は記入欄と紛らわしいので潰す
                    c = _span(t, r, used, COLS - 1)
                    _cell_text(c, ""); _shade(c)
                t.rows[r].height = Cm(0.8)
            elif kind == "word":
                for j, lab in enumerate(labs):
                    base = j * 5
                    _cell_text(t.cell(r, base), lab, 8, True, WD_ALIGN_PARAGRAPH.CENTER)
                    _cell_text(_span(t, r, base + 1, base + 4), "", 11, en=True)
                if len(labs) == 1:
                    c = _span(t, r, 5, 9); _cell_text(c, ""); _shade(c)
                t.rows[r].height = Cm(0.85)
            else:
                _cell_text(t.cell(r, 0), labs[0], 8, True, WD_ALIGN_PARAGRAPH.CENTER)
                _cell_text(_span(t, r, 1, 9), "", 11, en=True)
                t.rows[r].height = Cm(0.95)
            r += 1
        para(doc, "", 3, after=2)

    _anchors(doc)
    doc.save(path)
    return total

def build_key_csv(exam, path):
    """採点用の対応表。スキャンした解答用紙と突き合わせて採点するときに使う。"""
    import csv
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        wcsv = csv.writer(f)
        wcsv.writerow(["大問", "コース", "問", "種別", "正解", "配点"])
        for sec, cname, it, kind in answer_fields(exam):
            ty = it.get("type")
            if ty == "mcq":
                ans = "{}（{}）".format(KANA[it.get("answer", 0)], strip((it.get("choices") or [""])[it.get("answer", 0)]))
            elif ty == "bankpick":
                ans = "{}（{}）".format(KANA[it.get("answer", 0)], strip((it.get("bank") or [""])[it.get("answer", 0)]))
            elif ty == "wordorder":
                ans = strip(it.get("display") or it.get("answer", ""))
            else:
                ans = " / ".join(strip(a) for a in (it.get("answers") or []))
            wcsv.writerow([sec["no"], cname, it.get("label", ""),
                           {"sel": "選択", "word": "記述(語)", "long": "記述(文)"}[kind], ans, it.get("pt", 0)])

if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit(__doc__)
    js, outdir = sys.argv[1], sys.argv[2]
    exam = load_exam(js)
    name = sys.argv[3] if len(sys.argv) > 3 else strip(exam.get("title", "test"))
    os.makedirs(outdir, exist_ok=True)
    qp = os.path.join(outdir, name + "_問題.docx")
    ap = os.path.join(outdir, name + "_解答.docx")
    sp = os.path.join(outdir, name + "_解答用紙.docx")
    kp = os.path.join(outdir, name + "_採点キー.csv")
    t1 = build_question(exam, qp, name)
    t2 = build_answer(exam, ap, name)
    t3 = build_sheet(exam, sp, name)
    build_key_csv(exam, kp)
    print("✓ {} （{}点満点）".format(qp, t1))
    print("✓ {} （{}点満点）".format(ap, t2))
    print("✓ {} （{}点満点・A4 1枚）".format(sp, t3))
    print("✓ {}".format(kp))
